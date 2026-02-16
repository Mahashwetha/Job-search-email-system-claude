"""
Resume Tailor - Generates per-company tailored resumes using Gemini Flash.

Usage:
  Batch (from tracker):  python resume_tailor.py
  Single job:            python resume_tailor.py "https://job-posting-url" "Company Name"

Reads the Excel tracker for applied companies (status=done) with non-LinkedIn
role links, fetches job descriptions, and uses Gemini to suggest minimal
resume tweaks (skill reordering, keyword additions from existing experience).

NEVER fabricates experience — only reorders and surfaces existing skills.
"""

import json
import os
import re
import shutil
import sys
import tempfile
import time
from datetime import datetime
import openpyxl
import requests
from bs4 import BeautifulSoup
from docx import Document

try:
    from config import (
        TRACKER_FILE, GOOGLE_API_KEY, BASE_RESUME_PATH, RESUME_OUTPUT_DIR,
    )
except ImportError:
    print("ERROR: config.py not found!")
    print("Please copy config.template.py to config.py and fill in your details.")
    exit(1)

# Gemini Flash endpoint (free tier)
GEMINI_MODEL = "gemini-2.5-flash"
GEMINI_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{GEMINI_MODEL}:generateContent"
)


# ── Excel reading ────────────────────────────────────────────────────────────

def read_done_companies():
    """Read tracker, return companies with status=done and non-LinkedIn role links.
    Returns: {company: [{'role': str, 'role_link': str}, ...]}
    """
    temp_file = None
    try:
        try:
            wb = openpyxl.load_workbook(TRACKER_FILE, data_only=False)
        except PermissionError:
            print("  File locked — reading from temp copy...")
            fd, temp_file = tempfile.mkstemp(suffix='.xlsx')
            os.close(fd)
            shutil.copy2(TRACKER_FILE, temp_file)
            wb = openpyxl.load_workbook(temp_file, data_only=False)

        ws = wb.active
        companies = {}

        for i, row in enumerate(ws.iter_rows()):
            if i == 0:
                continue

            company = str(row[0].value or '').strip()
            role = str(row[1].value or '').strip()
            role_link = row[2].value if len(row) > 2 else None
            if not role_link and len(row) > 2 and row[2].hyperlink:
                role_link = row[2].hyperlink.target
            status = str(row[3].value or '').strip()

            if not company or 'Program/Product' in company:
                continue
            if 'done' not in status.lower():
                continue

            role_link_str = str(role_link).strip() if role_link else ''
            if not role_link_str.startswith('http'):
                role_link_str = ''

            # Skip LinkedIn links (can't scrape)
            if role_link_str and 'linkedin.com' in role_link_str.lower():
                continue

            if not role_link_str:
                continue

            if company not in companies:
                companies[company] = []

            if role and role != 'None':
                companies[company].append({
                    'role': role,
                    'role_link': role_link_str,
                })

        wb.close()
        if temp_file and os.path.exists(temp_file):
            os.remove(temp_file)

        print(f"Found {len(companies)} applied companies with fetchable role links")
        return companies
    except Exception as e:
        if temp_file and os.path.exists(temp_file):
            os.remove(temp_file)
        print(f"Error reading tracker: {e}")
        return {}


# ── Resume reading ───────────────────────────────────────────────────────────

def extract_resume_text(doc):
    """Extract paragraph texts with indices from a python-docx Document.
    Returns list of {'index': int, 'style': str, 'text': str}.
    """
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        paragraphs.append({
            'index': i,
            'style': p.style.name if p.style else 'Normal',
            'text': p.text.strip(),
        })
    return paragraphs


def build_resume_summary(paragraphs):
    """Build a human-readable resume summary for the LLM prompt."""
    lines = []
    for p in paragraphs:
        if p['text']:
            lines.append(f"[{p['index']}] ({p['style']}) {p['text']}")
    return '\n'.join(lines)


# ── Job description fetching ─────────────────────────────────────────────────

def fetch_job_description(url):
    """Fetch and extract text from a job posting URL.
    Tries JSON-LD structured data first (works for Workday/ATS sites),
    then falls back to plain text extraction."""
    try:
        headers = {
            'User-Agent': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/120.0.0.0 Safari/537.36'
            )
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, 'html.parser')

        # Try JSON-LD first (Workday, Greenhouse, etc. embed job data here)
        ld_script = soup.find('script', type='application/ld+json')
        if ld_script and ld_script.string:
            try:
                ld_data = json.loads(ld_script.string)
                desc = ld_data.get('description', '')
                title = ld_data.get('title', '')
                if desc and len(desc) > 100:
                    # Clean HTML from description
                    desc_soup = BeautifulSoup(desc, 'html.parser')
                    text = desc_soup.get_text(separator='\n', strip=True)
                    if title:
                        text = f"Job Title: {title}\n\n{text}"
                    print(f"    Extracted from JSON-LD ({len(text)} chars)")
                    if len(text) > 8000:
                        text = text[:8000] + '\n... [truncated]'
                    return text
            except (json.JSONDecodeError, TypeError):
                pass

        # Fallback: plain text extraction
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()

        text = soup.get_text(separator='\n', strip=True)

        if len(text) > 8000:
            text = text[:8000] + '\n... [truncated]'

        return text
    except Exception as e:
        print(f"    Failed to fetch {url}: {e}")
        return None


# ── Gemini LLM call ──────────────────────────────────────────────────────────

TAILOR_PROMPT = """You are a resume tailoring assistant. Given a candidate's resume and a job posting, suggest MINIMAL changes to better match the job.

STRICT RULES:
- ONLY reorder existing skills to prioritize job-relevant ones
- ONLY add keywords to bullet points if the candidate genuinely did that work (based on their existing experience)
- NEVER invent new skills, technologies, domains, or experience
- NEVER fabricate or exaggerate — only surface and reword existing experience
- Keep the resume to fit the same space (don't make text longer)
- Keep the original job title heading "SOFTWARE DEVELOPER SENIOR SPECIALIST" unchanged
- Remove the line "Open Minded to learn new technologies and languages." (paragraph 10)
- Keep "Open to remote, hybrid, or on-site roles..." line as-is

Return ONLY valid JSON with keys: skills_reorder (array of strings), profile_tagline (string), bullet_tweaks (array of objects with index/original/new), summary_tweak (string).

For skills_reorder: list ALL existing skill lines (paragraphs 24-31) in a new order that prioritizes skills relevant to this job. Use the exact existing text — do not add new skills.

For bullet_tweaks: only include bullets that benefit from minor keyword additions drawn from the candidate's ACTUAL experience. Most bullets should stay unchanged. Only tweak 3 bullets max. Keep each "new" text concise.

For profile_tagline: adjust the tagline keywords to better match the job, using ONLY domains the candidate actually has experience in.

For summary_tweak: lightly adjust the summary (paragraph 57) to emphasize aspects relevant to this job, keeping it the same length.

RESUME:
{resume_text}

JOB POSTING:
{job_description}
"""


def _repair_json(text):
    """Try to repair truncated JSON by closing open brackets/braces."""
    # Common case: Gemini output cut off mid-array or mid-object
    # Strategy: try progressively adding closing tokens
    for suffix in [']}\n}', ']\n}', '"}]\n}', '"}\n]\n}', '\n}']:
        try:
            return json.loads(text + suffix)
        except json.JSONDecodeError:
            continue
    return None


def call_gemini(resume_text, job_description, max_retries=3):
    """Call Gemini Flash to get tailoring suggestions with retry on rate limit."""
    # Use string concatenation instead of .format() to avoid issues with
    # braces in resume_text or job_description
    prompt = TAILOR_PROMPT.replace('{resume_text}', resume_text).replace('{job_description}', job_description)

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.2,
            "maxOutputTokens": 8192,
            "responseMimeType": "application/json",
        },
    }

    for attempt in range(max_retries):
        resp = requests.post(
            f"{GEMINI_URL}?key={GOOGLE_API_KEY}",
            json=payload,
            timeout=90,
        )

        if resp.status_code == 429:
            wait = 10 * (attempt + 1)
            print(f"    Rate limited, waiting {wait}s (attempt {attempt + 1}/{max_retries})...")
            time.sleep(wait)
            continue

        resp.raise_for_status()

        data = resp.json()
        candidates = data.get('candidates', [])
        if not candidates:
            raise ValueError(f"No candidates in Gemini response: {json.dumps(data)[:500]}")

        text = candidates[0]['content']['parts'][0]['text']

        # Strip markdown fences if present
        text = text.strip()
        if text.startswith('```'):
            text = re.sub(r'^```(?:json)?\s*', '', text)
            text = re.sub(r'\s*```$', '', text)

        # Extract JSON object if surrounded by other text
        json_match = re.search(r'\{[\s\S]*\}', text)
        if json_match:
            text = json_match.group(0)

        try:
            result = json.loads(text)
        except json.JSONDecodeError as e:
            # Try to repair truncated JSON (e.g. missing closing brackets)
            repaired = _repair_json(text)
            if repaired:
                result = repaired
            else:
                print(f"    JSON parse error: {e}")
                print(f"    Raw text (first 500): {text[:500]}")
                raise

        # Validate expected keys
        if not isinstance(result, dict):
            raise ValueError(f"Expected dict, got {type(result)}: {str(result)[:200]}")

        return result

    raise Exception("Gemini rate limit exceeded after all retries")


# ── DOCX manipulation ────────────────────────────────────────────────────────

def apply_tailoring(base_path, output_path, changes):
    """Copy base resume DOCX and apply tailoring changes.
    Returns a list of diff strings describing what changed.
    """
    doc = Document(base_path)
    paragraphs = doc.paragraphs
    diffs = []

    # 1. Remove "Open Minded" line (paragraph 10)
    if len(paragraphs) > 10 and 'open minded' in paragraphs[10].text.lower():
        paragraphs[10].text = ''
        diffs.append('Removed: "Open Minded to learn new technologies and languages."')

    # 2. Update profile tagline (paragraph 8)
    tagline = changes.get('profile_tagline', '').strip()
    if tagline and len(paragraphs) > 8:
        old = paragraphs[8].text.strip()
        if old != tagline:
            _replace_paragraph_text(paragraphs[8], tagline)
            diffs.append(f'Tagline: "{old}" -> "{tagline}"')

    # 3. Reorder skills (paragraphs 24-31)
    skills_reorder = changes.get('skills_reorder', [])
    if skills_reorder:
        skill_indices = list(range(24, 32))
        old_skills = [paragraphs[i].text.strip() for i in skill_indices if i < len(paragraphs)]
        available = min(len(skills_reorder), len(skill_indices))
        for j in range(available):
            idx = skill_indices[j]
            if idx < len(paragraphs):
                _replace_paragraph_text(paragraphs[idx], skills_reorder[j])
        if old_skills != skills_reorder[:len(old_skills)]:
            diffs.append(f'Skills reordered: {" | ".join(s.split(":")[0].strip() for s in skills_reorder)}')

    # 4. Tweak summary (paragraph 57)
    summary = changes.get('summary_tweak', '').strip()
    if summary and len(paragraphs) > 57:
        old = paragraphs[57].text.strip()
        if old != summary:
            _replace_paragraph_text(paragraphs[57], summary)
            diffs.append(f'Summary tweaked (paragraph 57)')

    # 5. Apply bullet tweaks
    for tweak in changes.get('bullet_tweaks', []):
        idx = tweak.get('index')
        new_text = tweak.get('new', '').strip()
        if idx is not None and new_text and idx < len(paragraphs):
            original = tweak.get('original', '')
            current = paragraphs[idx].text.strip()
            if original and _text_similar(current, original):
                _replace_paragraph_text(paragraphs[idx], new_text)
                diffs.append(f'Bullet [{idx}]: "{current[:60]}..." -> "{new_text[:60]}..."')
            else:
                print(f"    Skipped bullet tweak at [{idx}]: text mismatch")

    doc.save(output_path)
    return diffs


def _strip_markdown(text):
    """Remove markdown bold/italic markers that Gemini sometimes adds."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text


def _replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving formatting of the first run."""
    new_text = _strip_markdown(new_text)
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Keep first run's formatting, set its text, clear the rest
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''


def _text_similar(a, b):
    """Check if two texts are roughly similar (first 40 chars match)."""
    a_clean = re.sub(r'\s+', ' ', a.strip().lower())[:40]
    b_clean = re.sub(r'\s+', ' ', b.strip().lower())[:40]
    return a_clean == b_clean


def safe_company_name(company):
    """Convert company name to filesystem-safe string."""
    return re.sub(r'[^a-z0-9_]', '_', company.lower().strip()).strip('_')


# ── Diff summary ─────────────────────────────────────────────────────────────

def print_diff_summary(company, diffs, output_path):
    """Print a readable bullet-point diff of orig vs tailored resume."""
    print(f"\n    --- Changes for {company} ---")
    if not diffs:
        print("    (no changes applied)")
    else:
        for d in diffs:
            print(f"    * {d}")
    print(f"    Output: {output_path}")


def write_summary_file(company, role_name, url, diffs, changes, output_path):
    """Write a markdown summary file alongside the tailored resume."""
    summary_path = os.path.join(
        os.path.dirname(output_path),
        f"summary_changes_resume_{safe_company_name(company)}.md",
    )
    lines = [
        f"# Resume Tailoring Summary — {company}",
        f"",
        f"**Role:** {role_name}",
        f"**Job URL:** {url}",
        f"**Resume:** {os.path.basename(output_path)}",
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"",
        f"## Changes Applied",
        f"",
    ]
    if not diffs:
        lines.append("_(no changes applied)_")
    else:
        for d in diffs:
            lines.append(f"- {d}")

    # Include raw Gemini suggestions for reference
    lines.append("")
    lines.append("## Gemini Suggestions (raw)")
    lines.append("")
    if changes.get('profile_tagline'):
        lines.append(f"**Tagline:** {changes['profile_tagline']}")
    if changes.get('skills_reorder'):
        lines.append(f"**Skills order:** {' | '.join(changes['skills_reorder'])}")
    if changes.get('summary_tweak'):
        lines.append(f"**Summary:** {changes['summary_tweak']}")
    if changes.get('bullet_tweaks'):
        lines.append("**Bullet tweaks:**")
        for tw in changes['bullet_tweaks']:
            lines.append(f"  - [{tw.get('index')}] \"{tw.get('original', '')[:80]}...\" → \"{tw.get('new', '')[:80]}...\"")

    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines) + '\n')
    print(f"    Summary: {summary_path}")


# ── Core tailoring logic (shared by batch + single) ─────────────────────────

def _validate_config():
    """Check config values are set."""
    if not os.path.exists(BASE_RESUME_PATH):
        print(f"ERROR: Base resume not found: {BASE_RESUME_PATH}")
        exit(1)
    if GOOGLE_API_KEY == 'your_gemini_api_key_here':
        print("ERROR: Set GOOGLE_API_KEY in config.py")
        print("Get a free key at: https://aistudio.google.com/apikey")
        exit(1)


def tailor_one(company, role_name, url, resume_text, output_path):
    """Tailor resume for a single company/job. Returns True on success."""
    print(f"    Fetching JD: {url}")
    jd_text = fetch_job_description(url)

    if not jd_text or len(jd_text.strip()) < 200:
        print(f"    FAILED: Could not fetch job description ({len(jd_text or '')} chars)")
        return False

    job_context = f"Company: {company}\nRole: {role_name}\n\n{jd_text}"

    print(f"    Calling Gemini ({GEMINI_MODEL})...")
    try:
        changes = call_gemini(resume_text, job_context)
    except Exception as e:
        print(f"    FAILED Gemini call: {e}")
        return False

    print(f"    Applying tailoring changes...")
    try:
        diffs = apply_tailoring(BASE_RESUME_PATH, output_path, changes)
        print_diff_summary(company, diffs, output_path)
        write_summary_file(company, role_name, url, diffs, changes, output_path)
        return True
    except Exception as e:
        print(f"    FAILED to apply changes: {e}")
        if os.path.exists(output_path):
            os.remove(output_path)
        return False


# ── Batch mode (from Excel tracker) ─────────────────────────────────────────

def run_tailor():
    """Batch mode: read tracker, tailor all applied companies.
    Called from daily_job_search.py or standalone."""
    print("\n--- Resume Tailor ---")
    _validate_config()
    os.makedirs(RESUME_OUTPUT_DIR, exist_ok=True)

    base_doc = Document(BASE_RESUME_PATH)
    resume_paragraphs = extract_resume_text(base_doc)
    resume_text = build_resume_summary(resume_paragraphs)

    companies = read_done_companies()
    if not companies:
        print("  No companies to process.")
        return

    generated = 0
    skipped = 0
    failed = 0

    first_call = True
    for company, roles in companies.items():
        safe_name = safe_company_name(company)
        output_path = os.path.join(RESUME_OUTPUT_DIR, f'resume_{safe_name}.docx')

        # Idempotent: skip if already exists
        if os.path.exists(output_path):
            print(f"  SKIP {company}: already exists")
            skipped += 1
            continue

        # Rate limit: wait between Gemini calls (free tier)
        if not first_call:
            time.sleep(5)
        first_call = False

        print(f"\n  Processing: {company}")

        # Try each role link until one works
        success = False
        for role_info in roles:
            role_name = role_info['role']
            url = role_info['role_link']
            success = tailor_one(company, role_name, url, resume_text, output_path)
            if success:
                break

        if success:
            generated += 1
        else:
            failed += 1

    print(f"\n  Resume Tailor: {generated} generated, {skipped} skipped, {failed} failed")
    print(f"  Output dir: {RESUME_OUTPUT_DIR}")


# ── Single mode (CLI with URL) ──────────────────────────────────────────────

def run_single(url, company):
    """Single mode: tailor resume for one job URL.
    Usage: python resume_tailor.py "https://job-url" "Company Name"
    """
    print(f"\n--- Resume Tailor (single) ---")
    _validate_config()
    os.makedirs(RESUME_OUTPUT_DIR, exist_ok=True)

    base_doc = Document(BASE_RESUME_PATH)
    resume_paragraphs = extract_resume_text(base_doc)
    resume_text = build_resume_summary(resume_paragraphs)

    safe_name = safe_company_name(company)
    output_path = os.path.join(RESUME_OUTPUT_DIR, f'resume_{safe_name}.docx')

    if os.path.exists(output_path):
        print(f"  File already exists: {output_path}")
        print(f"  Delete it first to regenerate.")
        return

    print(f"  Company: {company}")
    success = tailor_one(company, company, url, resume_text, output_path)
    if success:
        print(f"\n  Done! Resume saved to: {output_path}")
        # Auto-generate outreach drafts
        try:
            from outreach_drafter import run_outreach
            run_outreach()
        except Exception as e:
            print(f"  Warning: outreach drafter failed: {e}")
    else:
        print(f"\n  Failed to generate resume for {company}")


# ── Entry point ──────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) >= 3:
        # Single mode: python resume_tailor.py "url" "Company"
        url = sys.argv[1]
        company = sys.argv[2]
        run_single(url, company)
    elif len(sys.argv) == 2 and sys.argv[1].startswith('http'):
        print("Usage: python resume_tailor.py \"https://job-url\" \"Company Name\"")
        print("Missing company name argument.")
        exit(1)
    else:
        # Batch mode from tracker
        run_tailor()


if __name__ == '__main__':
    main()
